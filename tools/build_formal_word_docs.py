from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "213B63"
BROWN = "743329"
GOLD = "D9A04B"
CORAL = "F15B55"
IVORY = "FAF6ED"
PALE = "F3F6FA"
MUTED = "627087"
LIGHT_GOLD = "FFF5E7"
WHITE = "FFFFFF"
BLACK = "252A34"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_cm[min(idx, len(widths_cm) - 1)]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep(paragraph, with_next: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together


def set_east_asia(run, font: str = "Microsoft YaHei") -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font)


def set_run_color(run, hex_color: str) -> None:
    run.font.color.rgb = RGBColor.from_string(hex_color)


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_inline(paragraph, text: str, default_bold: bool = False) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = default_bold
            set_east_asia(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, r_fonts)
            r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            set_run_color(run, BROWN)
        set_east_asia(run, run.font.name or "Microsoft YaHei")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = default_bold
        set_east_asia(run)


def configure_styles(doc: Document, compact: bool) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(8.5 if compact else 9.3)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = normal.paragraph_format
    pf.space_after = Pt(2.5 if compact else 3.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.05 if compact else 1.14
    pf.widow_control = True

    specs = {
        "Title": (25 if compact else 30, NAVY, 0, 7),
        "Subtitle": (13 if compact else 15, BROWN, 0, 8),
        "Heading 1": (14.5 if compact else 16.5, NAVY, 2, 5),
        "Heading 2": (11.5 if compact else 12.5, BROWN, 2, 3),
        "Heading 3": (9.8 if compact else 10.5, NAVY, 1, 2),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Figure Caption" not in doc.styles:
        cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Figure Caption"]
    cap.font.name = "Microsoft YaHei"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cap.font.size = Pt(7.6 if compact else 8.1)
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(3 if compact else 5)
    cap.paragraph_format.keep_together = True

    if "Lead" not in doc.styles:
        lead = doc.styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = doc.styles["Lead"]
    lead.font.name = "Microsoft YaHei"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lead.font.size = Pt(10 if compact else 10.5)
    lead.font.color.rgb = RGBColor.from_string(BROWN)
    lead.paragraph_format.space_before = Pt(4)
    lead.paragraph_format.space_after = Pt(5)
    lead.paragraph_format.left_indent = Cm(0.35)
    lead.paragraph_format.right_indent = Cm(0.35)

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(8.4 if compact else 9.1)
        style.paragraph_format.left_indent = Cm(0.7)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(1.5 if compact else 2.5)


def configure_page(doc: Document, compact: bool, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45 if compact else 1.7)
    section.bottom_margin = Cm(1.35 if compact else 1.55)
    section.left_margin = Cm(1.45 if compact else 1.65)
    section.right_margin = Cm(1.45 if compact else 1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.75)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"祠语智游  ·  {title}")
    set_east_asia(run)
    run.font.size = Pt(7.2)
    set_run_color(run, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("—  ")
    set_east_asia(run)
    run.font.size = Pt(7.2)
    set_run_color(run, MUTED)
    add_field(fp, "PAGE")
    run = fp.add_run("  —")
    set_east_asia(run)
    run.font.size = Pt(7.2)
    set_run_color(run, MUTED)


def add_page_accent(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("◆  非遗文化 × 智能体 × 场馆服务  ◆")
    set_east_asia(run)
    run.font.size = Pt(8.5)
    run.font.bold = True
    set_run_color(run, GOLD)


def add_cover(doc: Document, lines: list[str]) -> int:
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    add_page_accent(doc)
    title = lines[0].lstrip("# ").strip()
    subtitle = lines[2].lstrip("# ").strip() if len(lines) > 2 else ""
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    add_inline(p, title)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, subtitle)
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(14)
    r = rule.add_run("━━━━━━━━━━━━━━━━━━━━")
    set_east_asia(r)
    set_run_color(r, GOLD)
    r.font.size = Pt(9)

    i = 3
    while i < len(lines) and "<!-- PAGEBREAK -->" not in lines[i]:
        raw = lines[i].strip()
        if raw:
            if raw.startswith(">"):
                p = doc.add_paragraph(style="Lead")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_shading(p, LIGHT_GOLD)
                add_inline(p, raw.lstrip("> "))
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, raw)
        i += 1
    return i + 1


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, i


def choose_widths(rows: list[list[str]], content_cm: float) -> list[float]:
    cols = max(len(r) for r in rows)
    scores = []
    for col in range(cols):
        max_len = max((len(r[col]) if col < len(r) else 0) for r in rows)
        scores.append(max(8, min(max_len, 36)))
    if cols == 2:
        scores = [max(scores[0], 12), max(scores[1], 24)]
    total = sum(scores)
    widths = [content_cm * s / total for s in scores]
    min_w = 3.0 if cols <= 3 else 2.2
    for idx, width in enumerate(widths):
        if width < min_w:
            delta = min_w - width
            widths[idx] = min_w
            largest = max(range(cols), key=lambda j: widths[j] if j != idx else -1)
            widths[largest] -= delta
    return widths


def add_table(doc: Document, rows: list[list[str]], content_cm: float, compact: bool) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = choose_widths(rows, content_cm)
    set_table_fixed(table, widths)
    for r_idx, source_row in enumerate(rows):
        row = table.rows[r_idx]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=45 if compact else 65, bottom=45 if compact else 65)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = source_row[c_idx] if c_idx < len(source_row) else ""
            add_inline(p, text, default_bold=r_idx == 0)
            for run in p.runs:
                run.font.size = Pt(7.2 if compact else 8.0)
                if r_idx == 0:
                    set_run_color(run, WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, image_path: Path, alt: str, content_width_cm: float, compact: bool) -> None:
    from PIL import Image

    with Image.open(image_path) as im:
        w, h = im.size
    max_w = content_width_cm
    max_h = 9.0 if compact else 12.7
    width = max_w
    height = width * h / w
    if height > max_h:
        height = max_h
        width = height * w / h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width), height=Cm(height))
    cap = doc.add_paragraph(style="Figure Caption")
    add_inline(cap, alt)


def add_heading(doc: Document, text: str, level: int, first_content_heading: bool = False):
    style = "Title" if level == 1 else f"Heading {min(level - 1, 3)}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p.paragraph_format.left_indent = Cm(0.05)
        if not first_content_heading:
            p.paragraph_format.space_before = Pt(2)
    add_inline(p, text)
    return p


def add_markdown_document(md_path: Path, out_path: Path, kind: str) -> None:
    compact = kind == "technical"
    doc = Document()
    configure_styles(doc, compact=compact)
    title = "技术架构说明" if compact else "命题赛道解决方案书"
    configure_page(doc, compact=compact, title=title)
    section = doc.sections[0]
    content_width_cm = 21 - section.left_margin.cm - section.right_margin.cm

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    pending_page_break = False
    if kind == "solution":
        i = add_cover(doc, lines)
        pending_page_break = True

    first_h2 = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "<!-- PAGEBREAK -->":
            pending_page_break = True
            first_h2 = True
            i += 1
            continue
        if stripped.startswith("!["):
            m = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
            if m:
                alt, rel = m.groups()
                add_image(doc, (md_path.parent / rel).resolve(), alt, content_width_cm, compact)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows, content_width_cm, compact)
            i = next_i
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = add_heading(doc, heading.group(2), level, first_content_heading=first_h2)
            if pending_page_break:
                p.paragraph_format.page_break_before = True
                pending_page_break = False
            if level == 2:
                first_h2 = False
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Lead")
            set_paragraph_shading(p, LIGHT_GOLD)
            add_inline(p, stripped.lstrip("> "))
            i += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+[.)]\s+", "", stripped))
            i += 1
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.55 if not compact else 0.45)
        p.paragraph_format.keep_together = False
        add_inline(p, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", required=True)
    args = parser.parse_args()
    root = Path(args.root)
    evaluation = root / "data" / "chen_clan_academy" / "evaluation"
    out_dir = root / "output" / "word"
    add_markdown_document(
        evaluation / "祠语智游_命题赛道解决方案书_正式版.md",
        out_dir / "祠语智游_命题赛道解决方案书_正式版.docx",
        "solution",
    )
    add_markdown_document(
        evaluation / "祠语智游_技术架构说明_正式版.md",
        out_dir / "祠语智游_技术架构说明_正式版.docx",
        "technical",
    )
    print(out_dir)


if __name__ == "__main__":
    main()
