from docx import Document
from docx.oxml.ns import qn

INPUT = r"D:\VScode_Project\codexspace\codex_agent\祠语智游_命题赛道解决方案书_第四部分完善版_交付.docx"
OUTPUT = r"D:\VScode_Project\codexspace\codex_agent\祠语智游_命题赛道解决方案书_第一部分优化版_交付.docx"


def set_runs(paragraph, parts):
    """Replace text while preserving the original paragraph-level formatting."""
    paragraph.clear()
    for text, bold in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "等线"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


doc = Document(INPUT)
paragraphs = doc.paragraphs
start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "一、项目摘要")
figure = next(i for i, p in enumerate(paragraphs) if i > start and p.text.strip() == "图 1 祠语智游完整文化旅程")
targets = [p for p in paragraphs[start + 1:figure] if p.text.strip()]
assert len(targets) == 3, f"Expected 3 summary paragraphs, got {len(targets)}"

set_runs(targets[0], [
    ("“祠语智游”是面向陈家祠等非遗场馆的", False),
    ("多角色沉浸式智能导游", True),
    ("，将游前规划、游中讲解与问答、游后回顾与延伸服务串联为可运行、可追踪、可回退的文化体验流程。它不是单一聊天入口，而是一名能理解游客需求、陪伴真实行程的数字导游。", False),
])
set_runs(targets[1], [
    ("游客可按", False),
    ("时长、兴趣、语言、讲解深度与表达风格", True),
    ("获得审核路线中的个性化游览方案；到站后，系统结合当前位置、已讲内容和18种受控风格提供点位讲解、连续追问、研究摘要与术语卡片，并以真实游览记录生成称号、文明打卡建议和审核周边推荐。", False),
])
set_runs(targets[2], [
    ("项目以", False),
    ("“确定性决策＋生成式表达”", True),
    ("实现创新与落地的统一：审核知识、空间关系、路线与状态由确定性系统控制，角色模型仅负责受约束表达。知识卡、路线模板、空间图与角色规则均可配置、可审核、可复用，使产品既具多元体验和实用价值，也具备向博物馆、古建景区及非遗场馆复制推广的能力。", False),
])

doc.save(OUTPUT)
print(OUTPUT)
